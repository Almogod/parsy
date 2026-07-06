"""
Parsy Backend — Level 2a: Fast Structural Parser
Handles digital PDFs, DOCX, HTML, TXT with high-throughput rule-based extraction.
Supports parallel page chunking for large documents.
"""
import io, re, asyncio, concurrent.futures, logging
from dataclasses import dataclass, field
from typing import Generator

import fitz                    # PyMuPDF
from docx import Document
from bs4 import BeautifulSoup
import chardet

from base_parser import BaseParser, CorruptFileError, UnsupportedFormatError

log = logging.getLogger("parsy.parsers.fast")


@dataclass
class ParsedBlock:
    block_type: str            # "heading" | "paragraph" | "table" | "list" | "code" | "hr"
    level: int                 # heading level (1-6), 0 for non-headings
    content: str
    page: int
    bbox: tuple | None = None  # (x0, y0, x1, y1)
    table_data: list[list[str]] = field(default_factory=list)


@dataclass
class FastParseResult:
    blocks: list[ParsedBlock]
    page_count: int
    raw_text: str
    tables: list[list[list[str]]]
    metadata: dict


# ── Heading heuristics ─────────────────────────────────────────────────────
def _classify_block(span_text: str, font_size: float, flags: int,
                    body_size: float, page: int) -> ParsedBlock:
    bold   = bool(flags & 2**4)
    italic = bool(flags & 2**1)
    size_ratio = font_size / max(body_size, 10)

    if size_ratio >= 1.8 or (size_ratio >= 1.4 and bold):
        level = 1
    elif size_ratio >= 1.4 or (size_ratio >= 1.2 and bold):
        level = 2
    elif size_ratio >= 1.2 or bold:
        level = 3
    else:
        level = 0

    btype = "heading" if level > 0 else "paragraph"
    return ParsedBlock(block_type=btype, level=level, content=span_text.strip(), page=page)


# ── PDF page parser (runs in thread pool) ─────────────────────────────────
def _parse_pdf_page(args: tuple) -> list[ParsedBlock]:
    page_bytes, page_num = args
    doc = fitz.open(stream=page_bytes, filetype="pdf")
    page = doc[0]
    blocks_out = []

    # Get body font size (mode of all spans)
    sizes = []
    raw_blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
    for b in raw_blocks:
        if b.get("type") == 0:
            for line in b.get("lines", []):
                for span in line.get("spans", []):
                    sizes.append(span.get("size", 12))
    body_size = sorted(sizes)[len(sizes)//2] if sizes else 12

    # Extract blocks with structure
    for b in raw_blocks:
        if b.get("type") == 0:  # text
            for line in b.get("lines", []):
                text = " ".join(s.get("text","") for s in line.get("spans",[]))
                if not text.strip(): continue
                span = line["spans"][0] if line["spans"] else {}
                pb = _classify_block(
                    text, span.get("size", 12), span.get("flags", 0),
                    body_size, page_num
                )
                pb.bbox = b.get("bbox")
                blocks_out.append(pb)

        elif b.get("type") == 1:  # image — mark for OCR sub-pipeline
            blocks_out.append(ParsedBlock(
                block_type="image", level=0,
                content="[IMAGE]", page=page_num, bbox=b.get("bbox")
            ))

    doc.close()
    return blocks_out


# ── Table extractor using pdfplumber ──────────────────────────────────────
def _extract_pdf_tables(data: bytes) -> list[list[list[str]]]:
    try:
        import pdfplumber
        tables = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                for tbl in page.extract_tables() or []:
                    cleaned = [[str(c or "").strip() for c in row] for row in tbl]
                    if cleaned:
                        tables.append(cleaned)
        return tables
    except Exception:
        return []


class FastStructuralParser(BaseParser):
    """
    Level 2a parser. Splits documents into page chunks and processes
    them in parallel using a ThreadPoolExecutor.
    """

    supported_extensions = frozenset({
        "pdf", "docx", "html", "htm", "txt", "rtf", "md",
    })

    def __init__(self, max_workers: int = 4):
        super().__init__()
        self.max_workers = max_workers

    # ── Dispatcher — delegates to _parse (called by BaseParser.parse) ──
    async def _parse(self, filename: str, data: bytes) -> FastParseResult:
        ext = filename.rsplit(".", 1)[-1].lower()

        if ext == "pdf":
            try:
                return await self._parse_pdf(data)
            except Exception as exc:
                raise CorruptFileError(
                    f"PDF parse failed: {exc}", filename=filename, cause=exc
                ) from exc
        elif ext == "docx":
            try:
                return self._parse_docx(data)
            except Exception as exc:
                raise CorruptFileError(
                    f"DOCX parse failed: {exc}", filename=filename, cause=exc
                ) from exc
        elif ext in ("html", "htm"):
            return self._parse_html(data)
        elif ext in ("txt", "rtf"):
            return self._parse_text(data)
        elif ext == "md":
            return self._parse_markdown(data)
        else:
            log.warning(
                "Unknown extension; falling back to plain text",
                extra={"filename": filename, "ext": ext},
            )
            return self._parse_text(data)

    # ── PDF: parallel page chunking ───────────────────────────────────
    async def _parse_pdf(self, data: bytes) -> FastParseResult:
        doc   = fitz.open(stream=data, filetype="pdf")
        pages = len(doc)

        # Serialize each page into its own 1-page PDF bytes
        page_payloads = []
        for i in range(pages):
            sub = fitz.open()
            sub.insert_pdf(doc, from_page=i, to_page=i)
            page_payloads.append((sub.tobytes(), i))
            sub.close()
        doc.close()

        # Run pages in parallel
        loop = asyncio.get_running_loop()
        with concurrent.futures.ThreadPoolExecutor(max_workers=self.max_workers) as pool:
            futures = [loop.run_in_executor(pool, _parse_pdf_page, p) for p in page_payloads]
            page_results = await asyncio.gather(*futures)

        all_blocks: list[ParsedBlock] = [b for page in page_results for b in page]

        # Extract tables separately (pdfplumber)
        tables = await loop.run_in_executor(None, _extract_pdf_tables, data)

        meta = await loop.run_in_executor(None, self._pdf_metadata, data)

        raw_text = "\n".join(b.content for b in all_blocks if b.block_type != "image")
        return FastParseResult(
            blocks=all_blocks, page_count=pages,
            raw_text=raw_text, tables=tables, metadata=meta
        )

    def _pdf_metadata(self, data: bytes) -> dict:
        doc = fitz.open(stream=data, filetype="pdf")
        m = doc.metadata or {}
        info = {
            "title":    m.get("title",""),
            "author":   m.get("author",""),
            "subject":  m.get("subject",""),
            "creator":  m.get("creator",""),
            "producer": m.get("producer",""),
            "createdAt":m.get("creationDate",""),
            "pageCount":len(doc),
        }
        doc.close()
        return {k:v for k,v in info.items() if v}

    # ── DOCX ──────────────────────────────────────────────────────────
    def _parse_docx(self, data: bytes) -> FastParseResult:
        doc    = Document(io.BytesIO(data))
        blocks = []
        tables = []
        styles_map = {
            "Heading 1":1,"Heading 2":2,"Heading 3":3,
            "Heading 4":4,"Heading 5":5,"Heading 6":6,
        }

        for i, para in enumerate(doc.paragraphs):
            txt = para.text.strip()
            if not txt: continue
            style = para.style.name if para.style else ""
            level = styles_map.get(style, 0)
            btype = "heading" if level > 0 else "paragraph"
            blocks.append(ParsedBlock(btype, level, txt, page=i//30))

        for tbl in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
            tables.append(rows)
            blocks.append(ParsedBlock("table", 0, "", page=0, table_data=rows))

        raw = "\n".join(b.content for b in blocks if b.block_type != "table")
        props = doc.core_properties
        meta = {
            "author":   props.author or "",
            "created":  str(props.created or ""),
            "modified": str(props.modified or ""),
            "title":    props.title or "",
        }
        return FastParseResult(blocks, len(blocks)//30 + 1, raw, tables, meta)

    # ── HTML ──────────────────────────────────────────────────────────
    def _parse_html(self, data: bytes) -> FastParseResult:
        enc = chardet.detect(data)["encoding"] or "utf-8"
        soup = BeautifulSoup(data.decode(enc, errors="replace"), "lxml")
        for tag in soup(["script","style","nav","footer","aside","header"]):
            tag.decompose()

        blocks = []
        tables = []
        HEADING_TAGS = {"h1":1,"h2":2,"h3":3,"h4":4,"h5":5,"h6":6}

        for el in soup.find_all(True):
            tag = el.name.lower()
            if tag in HEADING_TAGS:
                blocks.append(ParsedBlock("heading", HEADING_TAGS[tag], el.get_text(" ",strip=True), 0))
            elif tag == "p":
                txt = el.get_text(" ", strip=True)
                if txt: blocks.append(ParsedBlock("paragraph", 0, txt, 0))
            elif tag in ("ul","ol"):
                items = [li.get_text(" ",strip=True) for li in el.find_all("li",recursive=False)]
                for it in items:
                    blocks.append(ParsedBlock("list", 0, it, 0))
            elif tag == "table":
                rows = []
                for tr in el.find_all("tr"):
                    cells = [td.get_text(" ",strip=True) for td in tr.find_all(["td","th"])]
                    if cells: rows.append(cells)
                if rows:
                    tables.append(rows)
                    blocks.append(ParsedBlock("table", 0, "", 0, table_data=rows))

        raw = "\n".join(b.content for b in blocks)
        title_tag = soup.find("title")
        meta = {"title": title_tag.string if title_tag else "",
                "description": (soup.find("meta",{"name":"description"}) or {}).get("content","")}
        return FastParseResult(blocks, 1, raw, tables, meta)

    # ── Plain text ────────────────────────────────────────────────────
    def _parse_text(self, data: bytes) -> FastParseResult:
        enc = chardet.detect(data)["encoding"] or "utf-8"
        text = data.decode(enc, errors="replace")
        blocks = []
        for line in text.splitlines():
            t = line.strip()
            if not t: continue
            # Simple heuristics for headings
            if len(t) < 80 and t.isupper() and len(t) > 3:
                blocks.append(ParsedBlock("heading", 1, t, 0))
            elif re.match(r"^[-\*\u2022]\s+", t):
                blocks.append(ParsedBlock("list", 0, t[2:].strip(), 0))
            else:
                blocks.append(ParsedBlock("paragraph", 0, t, 0))
        return FastParseResult(blocks, 1, text, [], {})

    # ── Markdown ──────────────────────────────────────────────────────
    def _parse_markdown(self, data: bytes) -> FastParseResult:
        text = data.decode("utf-8", errors="replace")
        blocks = []
        tables = []
        table_buf = []

        for line in text.splitlines():
            t = line.rstrip()
            hm = re.match(r"^(#{1,6})\s+(.*)", t)
            if hm:
                blocks.append(ParsedBlock("heading", len(hm.group(1)), hm.group(2), 0))
                continue
            if "|" in t:
                table_buf.append([c.strip() for c in t.split("|") if c.strip()])
                continue
            if table_buf:
                # flush table (skip separator row)
                rows = [r for r in table_buf if not all(re.match(r"^[-:]+$",c) for c in r)]
                if rows:
                    tables.append(rows)
                    blocks.append(ParsedBlock("table",0,"",0,table_data=rows))
                table_buf = []
            if re.match(r"^\s*[-\*]\s+", t):
                blocks.append(ParsedBlock("list",0,t.lstrip("-* "),0))
            elif t.strip():
                blocks.append(ParsedBlock("paragraph",0,t,0))

        return FastParseResult(blocks, 1, text, tables, {})
