"""
Parsy Backend — Level 1: Document Router
Inspects incoming files in milliseconds and selects execution path.
"""
import io, fitz, struct
from dataclasses import dataclass
from enum import Enum
from typing import BinaryIO


class Route(str, Enum):
    FAST_TEXT   = "fast_text"    # Digital PDF / DOCX / plain text
    VISION_OCR  = "vision_ocr"   # Scanned images, rotated pages
    STRUCTURED  = "structured"   # Spreadsheets, CSV, JSON, XML
    MARKDOWN    = "markdown"     # MD files — pass-through + enhance


@dataclass
class RouterDecision:
    route: Route
    confidence: float          # 0–1
    reasons: list[str]
    page_count: int
    has_text_layer: bool
    has_images: bool
    estimated_complexity: str  # "low" | "medium" | "high"
    recommended_workers: int


class DocumentRouter:
    """
    Inspects a file's byte signature, internal structure, and heuristics
    to route it to the cheapest/fastest parsing pipeline.
    """

    TEXT_DENSITY_THRESHOLD = 0.15   # chars/byte ratio for a "text-rich" PDF
    MAX_FAST_PAGES         = 200    # pages above this split into parallel chunks

    def route(self, filename: str, data: bytes) -> RouterDecision:
        ext = filename.rsplit(".", 1)[-1].lower()

        if ext in ("csv", "json", "xml", "xlsx", "xls", "ods"):
            return self._structured(data, ext)
        if ext == "md":
            return self._markdown(data)
        if ext in ("txt", "rtf"):
            return self._fast_text(data, ext, reasons=["Plain text — direct extraction"])
        if ext in ("html", "htm"):
            return self._fast_text(data, ext, reasons=["HTML — strip tags + structure"])
        if ext == "docx":
            return self._inspect_docx(data)
        if ext == "pdf":
            return self._inspect_pdf(data)

        return RouterDecision(
            route=Route.FAST_TEXT, confidence=0.5, reasons=["Unknown ext — fallback text"],
            page_count=1, has_text_layer=True, has_images=False,
            estimated_complexity="low", recommended_workers=1
        )

    # ── PDF inspection ────────────────────────────────────────────────
    def _inspect_pdf(self, data: bytes) -> RouterDecision:
        reasons = []
        try:
            doc = fitz.open(stream=data, filetype="pdf")
        except Exception as e:
            return RouterDecision(
                route=Route.VISION_OCR, confidence=0.9,
                reasons=[f"PDF corrupt/encrypted — OCR fallback: {e}"],
                page_count=0, has_text_layer=False, has_images=True,
                estimated_complexity="high", recommended_workers=4
            )

        page_count = len(doc)
        total_chars = 0
        image_pages = 0  # pages with NO usable text at all
        rotated = 0

        for page in doc:
            text = page.get_text()
            total_chars += len(text)
            images = page.get_images()
            # A page is truly "image-only" if it has virtually no text (<50 chars)
            # AND at least one image. A page with text + a logo is NOT an image page.
            if len(text.strip()) < 50 and len(images) > 0:
                image_pages += 1
            if page.rotation not in (0, 360):
                rotated += 1

        doc.close()

        text_density   = total_chars / max(len(data), 1)
        image_ratio    = image_pages / max(page_count, 1)
        # A PDF is text-rich if density is above threshold OR it has substantial total chars
        has_text_layer = text_density > self.TEXT_DENSITY_THRESHOLD or total_chars > 5000
        has_images     = image_pages > 0

        # Decision logic: text-layer quality takes priority over image presence
        if has_text_layer:
            route = Route.FAST_TEXT
            reasons.append(f"Text layer present ({total_chars:,} chars)")
            reasons.append(f"Text density: {text_density:.3f} chars/byte")
            if image_pages > 0:
                reasons.append(f"{image_pages} pages have images (ignored — text layer is dominant)")
            confidence = 0.95
        elif image_ratio > 0.5 or rotated > 0:
            route = Route.VISION_OCR
            reasons.append(f"{image_ratio:.0%} pages are scanned images")
            if rotated: reasons.append(f"{rotated} rotated page(s) detected")
            confidence = 0.9
        else:
            route = Route.VISION_OCR
            reasons.append("No reliable text layer — routing to OCR")
            confidence = 0.8

        complexity = (
            "high" if page_count > 100 or image_ratio > 0.3
            else "medium" if page_count > 20
            else "low"
        )
        workers = min(max(page_count // 10, 1), 8)

        return RouterDecision(
            route=route, confidence=confidence, reasons=reasons,
            page_count=page_count, has_text_layer=has_text_layer,
            has_images=has_images, estimated_complexity=complexity,
            recommended_workers=workers
        )

    # ── DOCX inspection ───────────────────────────────────────────────
    def _inspect_docx(self, data: bytes) -> RouterDecision:
        from docx import Document
        try:
            doc = Document(io.BytesIO(data))
            para_count  = len(doc.paragraphs)
            table_count = len(doc.tables)
            has_images  = any(r.element.tag.endswith('}drawing')
                              for p in doc.paragraphs for r in p.runs)
        except Exception:
            return self._fast_text(data, "docx", reasons=["DOCX fallback text mode"])

        return RouterDecision(
            route=Route.FAST_TEXT, confidence=0.95,
            reasons=[f"DOCX with {para_count} paragraphs, {table_count} tables"],
            page_count=max(para_count // 30, 1), has_text_layer=True,
            has_images=has_images, estimated_complexity="medium" if table_count > 5 else "low",
            recommended_workers=max(1, table_count // 3)
        )

    def _fast_text(self, data, ext, reasons):
        return RouterDecision(
            route=Route.FAST_TEXT, confidence=0.99, reasons=reasons,
            page_count=1, has_text_layer=True, has_images=False,
            estimated_complexity="low", recommended_workers=1
        )

    def _structured(self, data, ext):
        return RouterDecision(
            route=Route.STRUCTURED, confidence=0.99,
            reasons=[f"{ext.upper()} — structured data parser"],
            page_count=1, has_text_layer=True, has_images=False,
            estimated_complexity="low", recommended_workers=1
        )

    def _markdown(self, data):
        return RouterDecision(
            route=Route.MARKDOWN, confidence=0.99,
            reasons=["Markdown — AST parse + re-render"],
            page_count=1, has_text_layer=True, has_images=False,
            estimated_complexity="low", recommended_workers=1
        )
